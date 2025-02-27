import json
import boto3
import matplotlib.pyplot as plt
import numpy as np
import os
from io import StringIO, BytesIO
from docx import Document
from docx.shared import Inches, Pt
from botocore.exceptions import ClientError
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import datetime


def lambda_handler(event, context):
    client = boto3.resource("dynamodb")
    cost_table = client.Table("CostData_THB")
    account_table = client.Table("Accounts_THB")
    accounts = account_table.scan()['Items']
    curr_date = datetime.date.today()
    curr_date = datetime.date(curr_date.year, curr_date.month, 1)
    
    for line in accounts:
        data = cost_table.get_item(Key={
            'AccountNr': line['AccountNr'],
            'Date': str(curr_date)
        })
        costs = data['Item']['data']

        # verify clean plot
        plt.clf()

        bargraph = BytesIO()
        New_Colors = ['#0089BA', '#00A7C5', '#00C2BB', '#48DAA2', '#A3ED84', '#F9F871', '#5FBEF2']
        plt.bar(range(len(costs)), list(costs.values()), align='center', color=New_Colors)
        plt.title('AWS Cost Data', fontsize=14)
        plt.xlabel('Services', fontsize=14)
        plt.ylabel('Cost', fontsize=14)
        plt.xticks(range(len(costs)), list(costs.keys()), rotation=45)
        plt.legend(loc="upper left")

        plt.savefig(bargraph, format='png', bbox_inches='tight')
        bargraph.seek(0)

        # piechart
        # Data to plot
        # verify clean plot

        plt.clf()
        piechart = BytesIO()
        labels = []
        sizes = []

        for x, y in costs.items():
            labels.append(x)
            sizes.append(y)

        # Plot
        plt.pie(sizes, labels=labels)

        plt.axis('equal')
        plt.savefig(piechart, format='png', bbox_inches='tight')
        piechart.seek(0)

        document = Document()
        document.add_heading('Cost Data Visualization', 0)
        p = document.add_paragraph('Report')
        document.add_picture(bargraph, height=Inches(3.641732), width=Inches(5.1889764))
        document.add_picture(piechart, height=Inches(3.641732), width=Inches(5.1889764))
        doc = BytesIO()
        document.save(doc)
        doc.seek(0)

        send_email(doc)


def send_email(doc):
    # emailData = event['Records'][0]["messageAttributes"]
    # print(emailData['fromName'])

    # Replace sender@example.com with your "From" address.
    # This address must be verified with Amazon SES.
    SENDER = "Cost Data Messenger" + " <projthbingen@gmail.com>"

    # Replace recipient@example.com with a "To" address. If your account
    # is still in the sandbox, this address must be verified.
    RECIPIENT = "jonas.trojahn@th-bingen.de"

    # print(emailData)

    """if 'toCCEmail' in emailData:
      CC_RECIPIENT = emailData['toCCEmail']['stringValue']
      DESTINATION={
                        'ToAddresses': [
                            RECIPIENT,
                        ],
                        'CcAddresses': [
                            CC_RECIPIENT,
                        ],
                        'BccAddresses': [
                            'marketingcc@example.com',
                        ]
                    }

    else:"""
    DESTINATION = {
        'ToAddresses': [
            RECIPIENT,
        ],
        'BccAddresses': [
            'projthbingen@gmail.com',
        ]
    }

    # Specify a configuration set. If you do not want to use a configuration
    # set, comment the following variable, and the
    # ConfigurationSetName=CONFIGURATION_SET argument below.
    # CONFIGURATION_SET = "ConfigSet"

    # If necessary, replace us-west-2 with the AWS Region you're using for Amazon SES.
    AWS_REGION = "us-east-1"

    # The subject line for the email.
    SUBJECT = "AWS Cost Report"

    # The full path to the file that will be attached to the email.
    # ATTACHMENT = ""

    # The email body for recipients with non-HTML email clients.
    # BODY_TEXT = ("Amazon SES Test (Python)\r\n"
    #             "This email was sent with Amazon SES using the "
    #             "AWS SDK for Python (Boto)."
    #            )

    # The HTML body of the email.
    BODY_TEXT = "Guten Tag,\nIm Anhang finden Sie den Kostenbericht für diesen Monat"
    # emailData['body']['stringValue']

    # The character encoding for the email.
    CHARSET = "UTF-8"

    # Create a new SES resource and specify a region.
    client = boto3.client('ses', region_name=AWS_REGION)

    # Create a multipart/mixed parent container.
    msg = MIMEMultipart('mixed')
    # Add subject, from and to lines.
    msg['Subject'] = SUBJECT
    msg['From'] = SENDER
    msg['To'] = RECIPIENT

    # Create a multipart/alternative child container.
    msg_body = MIMEMultipart('alternative')

    # Encode the text and HTML content and set the character encoding. This step is
    # necessary if you're sending a message with characters outside the ASCII range.
    textpart = MIMEText(BODY_TEXT.encode(CHARSET), 'plain', CHARSET)
    # htmlpart = MIMEText(BODY_HTML.encode(CHARSET), 'html', CHARSET)

    # Add the text and HTML parts to the child container.
    msg_body.attach(textpart)
    # msg_body.attach(htmlpart)

    # Define the attachment part and encode it using MIMEApplication.
    att = MIMEApplication(doc.read())

    # Add a header to tell the email client to treat this part as an attachment,
    # and to give the attachment a name.
    att.add_header('Content-Disposition', 'attachment', filename='test.docx')

    # Attach the multipart/alternative child container to the multipart/mixed
    # parent container.
    msg.attach(msg_body)

    # Add the attachment to the parent container.
    msg.attach(att)
    # print(msg)

    # Try to send the email.
    try:
        # Provide the contents of the email.
        response = client.send_raw_email(
            Source=SENDER,
            Destinations=[
                RECIPIENT
            ],
            RawMessage={
                'Data': msg.as_string(),
            },
        )

    # Display an error if something goes wrong.
    except ClientError as e:
        print(e.response['Error']['Message'])
    else:
        print("Email sent! Message ID:")
        print(response['MessageId'])
        print(RECIPIENT)
